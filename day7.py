from docx import Document

def create_cv(name, email, phone, skills, experience, education):
    document = Document()

    document.add_heading('Curriculum Vitae', 0)
    document.add_paragraph('Name: {}'.format(name))
    document.add_paragraph('Email: {}'.format(email))
    document.add_paragraph('Phone: {}'.format(phone))
    document.add_page_break()

    document.add_heading('Skills', level=1)
    for skill in skills:
        document.add_paragraph(skill)

    document.add_page_break()

  
    document.add_heading('Experience', level=1)
    for exp in experience:
        document.add_paragraph(exp['title'], style='Heading 2')
        document.add_paragraph('{} - {}'.format(exp['start_date'], exp['end_date']))
        document.add_paragraph(exp['description'])

    document.add_page_break()

    document.add_heading('Education', level=1)
    for edu in education:
        document.add_paragraph(edu['degree'], style='Heading 2')
        document.add_paragraph('{} - {}'.format(edu['start_date'], edu['end_date']))
        document.add_paragraph(edu['university'])

    document.save('cv.docx')

if __name__ == '__main__':

    name = "Aayush Stha"
    email = "stha7009@gmail.com"
    phone = "980000000"
    skills = ["HTML", "CSS", "Java", "Python"]
    experience = [
        {
            'title': 'Python Intern',
            'start_date': 'November 2023',
            'end_date': 'February 2024',
            'description': 'Worked on various Python projects for 3 months.'
        }
    ]
    education = [
        {
            'degree': 'Bachelor of Computer Applications (BCA)',
            'start_date': '2022',
            'end_date': 'Present',
            'university': 'Itahari Namuna college'
        }
    ]

    create_cv(name, email, phone, skills, experience, education)
